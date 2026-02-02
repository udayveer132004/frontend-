"""Text extraction utilities for PDF and DOCX files."""

import os
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text content as a string
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file is not a valid PDF
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    text_content = []
    
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    return "\n\n".join(text_content)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content as a string
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file is not a valid DOCX
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        doc = Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        return "\n".join(text_content)
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text(file_path: str) -> str:
    """
    Extract text from a resume file (PDF or DOCX).
    
    Automatically detects the file type and uses the appropriate extraction method.
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Extracted text content as a string
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file format is not supported or extraction fails
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_extension = Path(file_path).suffix.lower()
    
    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file format: {file_extension}. "
            "Supported formats: .pdf, .docx"
        )
